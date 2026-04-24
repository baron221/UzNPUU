import os
from pypdf import PdfReader
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def load_knowledge_base(folder=None, include_files=None):
    if folder is None:
        folder = os.path.join(BASE_DIR, "knowledge")

    all_text = []

    if not os.path.exists(folder):
        print(f"Warning: Folder '{folder}' not found. Creating it...")
        os.makedirs(folder, exist_ok=True)
        return ""

    files = os.listdir(folder)
    if include_files is not None:
        files = [f for f in files if f in include_files]

    if not files:
        print(f"Warning: No files found/included in '{folder}'.")
        return ""

    for filename in files:
        path = os.path.join(folder, filename)
        text = ""
        try:
            if filename.lower().endswith(".pdf"):
                reader = PdfReader(path)
                text = "\n".join(p.extract_text() for p in reader.pages if p.extract_text())
                print(f"Loaded PDF: {filename}")
            elif filename.lower().endswith(".docx"):
                doc = Document(path)
                paras = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if len(cells) >= 2:
                            paras.append(f"Savol: {cells[0]}\nJavob: {cells[1]}")
                        elif len(cells) == 1:
                            # If a cell contains 'Savol:' or 'Javob:', append as is.
                            # Otherwise just append it, relying on parser to handle chunks.
                            paras.append(cells[0])
                text = "\n".join(paras)
                print(f"Loaded DOCX: {filename}")
            elif filename.lower().endswith((".txt", ".md")):
                with open(path, "r", encoding="utf-8") as f:
                    text = f.read()
                print(f"Loaded TXT: {filename}")
            elif filename.lower().endswith((".xlsx", ".xls")):
                import openpyxl
                wb = openpyxl.load_workbook(path)
                rows = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    rows.append(f"-- Sheet: {sheet} --")
                    for row in ws.iter_rows(values_only=True):
                        row_text = " | ".join(str(c) for c in row if c is not None)
                        if row_text.strip():
                            rows.append(row_text)
                text = "\n".join(rows)
                print(f"Loaded Excel: {filename}")
            else:
                print(f"Skipped: {filename}")
                continue

            if text.strip():
                all_text.append(f"=== FILE: {filename} ===\n{text.strip()}")
        except Exception as e:
            print(f"Error reading {filename}: {e}")

    print(f"\nKnowledge base: Loaded {len(all_text)} file(s).\n")
    return "\n\n".join(all_text)